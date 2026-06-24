#!/usr/bin/env python3
"""Branded Word document template (also imports cleanly into Google Docs).

A titled cover with the mark, in-brand headings (Space Grotesk), body copy, and a
color table with shaded swatch cells. Text uses Space Grotesk (renders if
installed; falls back gracefully).

CLI: docx_gen.py <out.docx> <icon.png>
"""
import sys
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Space Grotesk"
INK, AMBER, SLATE = RGBColor(0x15, 0x16, 0x1A), RGBColor(0xB5, 0x78, 0x1A), RGBColor(0x4A, 0x56, 0x5A)

def _run(p, text, size, color, bold=False):
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = bold
    return r

def _para(doc, text, size, color, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph(); p.alignment = align; p.paragraph_format.space_after = Pt(after)
    _run(p, text, size, color, bold)
    return p

def _shade(cell, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)

def build(out, icon):
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = FONT; normal.size = Pt(11); normal.color.rgb = INK

    # cover
    cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.add_run().add_picture(icon, width=Inches(1.4))
    _para(doc, "fastverk", 30, INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    _para(doc, "Brand — document template", 13, SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    _para(doc, "The mark", 17, AMBER, bold=True, after=4)
    _para(doc, "The mark plays off F, V, the down-triangle (forall) and the horizontal arms "
               "(exists) — a deliberate golden-ratio construction on the unit circle, not a "
               "hand-traced drawing. Every asset — this document included — derives from one "
               "parametric source.", 11, INK, after=14)

    _para(doc, "Color", 17, AMBER, bold=True, after=6)
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
    for cell, (hexc, name) in zip(table.rows[0].cells,
                                  [("15161A", "bg midnight"), ("ECE7DA", "fg cream"),
                                   ("E0A33E", "accent amber"), ("4A565A", "tertiary slate")]):
        _shade(cell, hexc)
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(18); para.paragraph_format.space_after = Pt(18)
        on = RGBColor(0xEC, 0xE7, 0xDA) if hexc in ("15161A", "4A565A") else INK
        _run(para, "#" + hexc, 10, on, bold=True)
        cap = cell.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cap, name, 8, on)

    doc.save(out)
    print("wrote", out)

if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
